"""Render already-paginated Depo-Pro testimony pages without deciding wrapping."""
from __future__ import annotations
import argparse, json
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree

def field(paragraph):
    run=paragraph.add_run()
    for kind,text in (("begin",None),(None," PAGE "),("separate",None),(None,"1"),("end",None)):
        node=OxmlElement("w:fldChar" if kind else ("w:instrText" if text==" PAGE " else "w:t"))
        if kind: node.set(qn("w:fldCharType"),kind)
        else: node.text=text
        run._r.append(node)

def set_font(run,profile):
    name=profile["font"]["family"];run.font.name=name;run.font.size=Pt(profile["font"]["pointSize"])
    fonts=run._element.get_or_add_rPr().get_or_add_rFonts();fonts.set(qn("w:ascii"),name);fonts.set(qn("w:hAnsi"),name)

def fit_administrative_line(run,line,page,profile):
    # Reviewed administrative templates occasionally express a rule as 68 underscore
    # characters.  They are document furniture, not testimony wrapping input.  Word otherwise
    # turns that one modeled physical line into two at the proven 63-cell testimony width and
    # spills the fixed page.  fitText preserves the approved template bytes and constrains the
    # run to the model's one physical line; testimony never enters this path.
    if page.get("sectionKind")!="administrative" or len(line.get("text", ""))<=profile["charactersPerLine"]:return
    fit=OxmlElement("w:fitText");fit.set(qn("w:val"),str(profile["text"]["widthTwips"]));run._element.get_or_add_rPr().append(fit)

def configure(section,profile):
    page=profile["page"];text=profile["text"];box=profile["formatBox"]
    section.page_width=Inches(page["widthTwips"]/1440);section.page_height=Inches(page["heightTwips"]/1440)
    section.left_margin=Inches(text["leftMarginTwips"]/1440);section.right_margin=Inches(text["rightMarginTwips"]/1440)
    section.top_margin=Inches(text["topMarginTwips"]/1440);section.bottom_margin=Inches(text["bottomMarginTwips"]/1440)
    section.header_distance=Inches(.15);section.footer_distance=Inches(.35)
    numbers=OxmlElement("w:lnNumType");numbers.set(qn("w:countBy"),"1");numbers.set(qn("w:start"),"0");numbers.set(qn("w:restart"),"newPage");numbers.set(qn("w:distance"),str(profile["lineNumbers"]["distanceTwips"]));section._sectPr.append(numbers)
    header=section.header;header.paragraphs[0].clear();run=header.paragraphs[0].add_run();pict=OxmlElement("w:pict");rect=etree.Element("{urn:schemas-microsoft-com:vml}rect")
    rect.set("id","DepoProTranscriptFormatBox");rect.set("filled","f");rect.set("stroked","t");rect.set("strokecolor","000000");rect.set("strokeweight",f'{box["borderPoints"]}pt')
    rect.set("style",f'position:absolute;margin-left:{box["leftInches"]*72:.3f}pt;margin-top:{box["topInches"]*72:.3f}pt;width:{box["widthInches"]*72:.3f}pt;height:{box["heightInches"]*72:.3f}pt;z-index:-251654144;mso-position-horizontal-relative:page;mso-position-vertical-relative:page')
    wrap=etree.Element("{urn:schemas-microsoft-com:office:word}wrap");wrap.set("type","none");wrap.set("anchorx","page");wrap.set("anchory","page");rect.append(wrap);pict.append(rect);run._r.append(pict)
    number=header.add_paragraph();number.alignment=WD_ALIGN_PARAGRAPH.RIGHT;number.paragraph_format.right_indent=Inches(box["rightClearanceInches"]);field(number)

def main():
    parser=argparse.ArgumentParser();parser.add_argument("--spec",required=True);parser.add_argument("--output",required=True);parser.add_argument("--mapping",required=True);args=parser.parse_args()
    spec=json.loads(Path(args.spec).read_text(encoding="utf-8"));profile=spec["profile"];pages=spec["pages"]
    if any(len(page["lines"])!=profile["linesPerPage"] for page in pages):raise ValueError("Every modeled page must contain the profile's physical line count")
    doc=Document();configure(doc.sections[0],profile);style=doc.styles["Normal"];style.font.name=profile["font"]["family"];style.font.size=Pt(profile["font"]["pointSize"]);style.paragraph_format.space_before=Pt(0);style.paragraph_format.space_after=Pt(0);style.paragraph_format.line_spacing=Pt(profile["text"]["lineSpacingTwips"]/20)
    doc.core_properties.subject=f'Depo-Pro fixed pages; model {spec.get("modelHash")}'
    mapping=[]
    for page_index,page in enumerate(pages):
        page_lines=list(page["lines"])
        if page_index==len(pages)-1:
            while page_lines and not page_lines[-1].get("occupied"):page_lines.pop()
        for line_index,line in enumerate(page_lines):
            paragraph=doc.add_paragraph();fmt=paragraph.paragraph_format;fmt.space_before=Pt(0);fmt.space_after=Pt(0);fmt.line_spacing=Pt(profile["text"]["lineSpacingTwips"]/20);fmt.widow_control=False;fmt.keep_together=True
            if page_index and line_index==0:fmt.page_break_before=True
            run=paragraph.add_run(line["text"]);set_font(run,profile);fit_administrative_line(run,line,page,profile)
            mapping.append({"modelPage":page["pageNumber"],"modelLine":line["position"],"docxParagraph":len(mapping)+1,"paragraphId":line.get("paragraphId"),"fragmentIds":line.get("fragmentIds",[]),"sourceWordIds":line.get("sourceWordIds",[])})
    output=Path(args.output).resolve();output.parent.mkdir(parents=True,exist_ok=True);doc.save(output);Path(args.mapping).write_text(json.dumps({"modelHash":spec.get("modelHash"),"profile":profile["id"],"lines":mapping},indent=2),encoding="utf-8")
    print(json.dumps({"renderer":"DEPO_PRO_INTERNAL_FIXED_PAGE_OOXML_V1","outputPath":str(output),"pages":len(pages),"physicalLines":len(mapping),"profile":profile["id"]}))
if __name__=="__main__":main()
